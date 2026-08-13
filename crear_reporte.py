from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path


OUT = Path(r"C:\Users\benav\OneDrive\Documents\almacen\Reporte_Tecnico_Venta_de_Libros.docx")

NAVY = "0B2545"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "5B6573"


def set_run_font(run, size=11, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_inches):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    table_pr.append(tbl_ind)
    for row in table.rows:
        for index, width in enumerate(widths_inches):
            cell = row.cells[index]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margin(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_body(doc, text, bold_start=None):
    p = doc.add_paragraph(style="Normal")
    if bold_start and text.startswith(bold_start):
        set_run_font(p.add_run(bold_start), bold=True)
        set_run_font(p.add_run(text[len(bold_start):]))
    else:
        set_run_font(p.add_run(text))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(text))
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run_font(p.add_run(header), size=10, bold=True, color=NAVY)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    header_mark = OxmlElement("w:tblHeader")
    header_mark.set(qn("w:val"), "true")
    tr_pr.append(header_mark)
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            p = cells[i].paragraphs[0]
            set_run_font(p.add_run(str(value)), size=9.5)
            if len(row_values) > 2 and i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    shade(cell, "F4F6F9")
    p = cell.paragraphs[0]
    set_run_font(p.add_run(title + " "), size=10.5, bold=True, color=BLUE_DARK)
    set_run_font(p.add_run(text), size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, BLUE_DARK, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.add_run("Venta de Libros Digital | Reporte técnico"), size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run("Guía de estudio y operación local"), size=8.5, color=MUTED)


def build():
    doc = Document()
    configure_styles(doc)

    # Editorial-cover opening, intentionally simple for a technical study guide.
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("REPORTE TÉCNICO"), size=12, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Venta de Libros Digital"), size=28, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Guía de estudio: backend, frontend, base de datos y pruebas"), size=14, color=BLUE_DARK)
    p.paragraph_format.space_after = Pt(26)
    add_callout(doc, "Objetivo.", "Explicar las correcciones realizadas y cómo comprobar que la aplicación funciona correctamente en tu equipo.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Proyecto local: React + Express + MySQL (XAMPP)"), size=10.5, color=MUTED)
    doc.add_page_break()

    add_heading(doc, "1. Resumen del proyecto")
    add_body(doc, "La aplicación se llama Venta de Libros Digital. Permite administrar la información de una librería: clientes, libros, ventas, pagos, usuarios y las demás tablas relacionadas.")
    add_body(doc, "El sistema se organizó en tres capas que se comunican por HTTP: el frontend muestra los módulos, el backend valida y procesa las solicitudes, y MySQL guarda los datos.")
    add_table(doc, ["Capa", "Tecnología", "Responsabilidad"], [
        ("Frontend", "React (puerto 3000)", "Pantallas, formularios, listado, edición y ocultar registros de la vista."),
        ("Backend", "Node.js + Express (puerto 5000)", "API REST, validación, CRUD y mensajes de error controlados."),
        ("Base de datos", "MySQL/XAMPP: libreria", "Tablas, claves primarias, claves foráneas y persistencia real."),
    ], [1.15, 1.9, 3.45])

    add_heading(doc, "2. Arquitectura y flujo de información")
    add_body(doc, "Flujo normal de una operación:")
    add_number(doc, "El usuario abre un módulo en http://localhost:3000/modulos/<modulo>.")
    add_number(doc, "React solicita los registros al backend, por ejemplo GET http://localhost:5000/api/cliente?limit=100.")
    add_number(doc, "Express ejecuta la consulta mediante el pool de MySQL sobre la base libreria.")
    add_number(doc, "La API responde JSON y React vuelve a dibujar la tabla.")
    add_number(doc, "Al guardar, React envía POST o PUT; al editar se actualiza la fila y se recarga la información.")
    add_callout(doc, "Direcciones clave.", "Frontend: http://localhost:3000. API: http://localhost:5000. Estado de la API: http://localhost:5000/api/health.")

    add_heading(doc, "3. Base de datos revisada")
    add_body(doc, "Se importó y configuró la base de datos libreria a partir del archivo SQL del proyecto. La aplicación usa las siguientes tablas:")
    add_table(doc, ["Grupo", "Tablas"], [
        ("Personas y acceso", "cliente, empleado, usuario, rol, rol_usuario"),
        ("Catálogo", "categoria, editorial, autor, libro, libro_autor"),
        ("Venta y cobro", "venta, detalle_venta, metodo_pago, pago"),
    ], [1.75, 4.75])
    add_body(doc, "Relaciones que el sistema respeta: un libro pertenece a una categoría y editorial; libro_autor une libros con autores; una venta requiere cliente y empleado; detalle_venta requiere venta y libro; pago requiere venta y método de pago; rol_usuario une usuarios con roles.")
    add_callout(doc, "Regla importante.", "No se puede guardar una relación con un ID que no exista. Por ejemplo, una venta no puede referenciar un cliente inexistente.")

    add_heading(doc, "4. Backend: cambios realizados")
    add_bullet(doc, "Se organizaron rutas y controladores para los módulos de las 14 tablas.")
    add_bullet(doc, "Se creó un controlador CRUD genérico que consulta las columnas reales de cada tabla y admite listar, crear, editar y eliminar cuando corresponde.")
    add_bullet(doc, "Se añadieron rutas especiales para las tablas con clave compuesta: libro_autor usa id_libro + id_autor y rol_usuario usa id_rol + id_usuario.")
    add_bullet(doc, "Se agregó GET /api/health para confirmar que Express está activo y conectado a MySQL.")
    add_bullet(doc, "Se mejoraron los mensajes de MySQL: conexión rechazada, registro relacionado inexistente y registro con relaciones activas.")
    add_bullet(doc, "La configuración local del backend apunta a DB_HOST=localhost, DB_USER=root, DB_NAME=libreria y PORT=5000.")
    add_table(doc, ["Ruta", "Uso"], [
        ("GET /api/health", "Comprueba servidor y conexión a la base de datos."),
        ("GET /api/<tabla>?limit=100", "Lista datos de una tabla."),
        ("POST /api/<tabla>", "Crea un registro válido."),
        ("PUT /api/<tabla>/<id>", "Actualiza un registro."),
        ("DELETE /api/<tabla>/<id>", "Elimina en la base de datos solo si no tiene relaciones."),
    ], [2.4, 4.1])

    add_heading(doc, "5. Frontend: cambios realizados")
    add_bullet(doc, "Se cambió el nombre visible de Almacén a Venta de Libros / Venta de Libros Digital, incluido título y favicon del navegador.")
    add_bullet(doc, "Se restauró el enfoque del frontend anterior y se aplicó un estilo claro, centrado y con tarjetas para los módulos.")
    add_bullet(doc, "Se creó una página genérica /modulos/:modulo para administrar cliente, libro, venta, pago y los demás módulos.")
    add_bullet(doc, "Los formularios cargan listas desplegables en los campos relacionados (por ejemplo cliente, empleado, libro o método de pago), evitando IDs inválidos.")
    add_bullet(doc, "Al guardar o editar, la tabla se actualiza con la respuesta de la API.")
    add_bullet(doc, "Se agregó almacenamiento local de instantáneas e historial para poder verlo en DevTools > Aplicación > Local Storage.")
    add_bullet(doc, "El botón antes llamado Eliminar se adaptó a Ocultar en la interfaz: quita el registro solamente de la página y no borra el dato en MySQL.")
    add_callout(doc, "Diferencia clave.", "Ocultar afecta únicamente la vista y se guarda en Local Storage. DELETE de la API sí intenta borrar de MySQL y puede ser rechazado si existen relaciones activas.")

    add_heading(doc, "6. Pruebas realizadas")
    add_table(doc, ["Prueba", "Resultado esperado", "Resultado"], [
        ("Estado de API", "GET /api/health responde OK", "Verificado: API y base de datos conectadas."),
        ("Lectura de módulos", "GET de las 14 tablas responde 200", "Verificado con límite de registros."),
        ("Claves compuestas", "Consulta de libro_autor y rol_usuario", "Verificado con sus dos IDs."),
        ("Guardar cliente", "POST crea un registro y aparece en la tabla", "Verificado desde la interfaz y persistido en MySQL."),
        ("Error de relación", "ID relacionado inexistente devuelve un mensaje claro", "Verificado: respuesta controlada de la API."),
        ("Compilación frontend", "React genera build sin errores", "Comprobación incluida en la revisión final."),
    ], [1.5, 2.75, 2.25])

    add_heading(doc, "7. Cómo ejecutar y comprobar el sistema")
    add_heading(doc, "7.1 Preparar servicios", level=2)
    add_number(doc, "Abre XAMPP y enciende MySQL. Si MySQL está detenido, el backend mostrará ECONNREFUSED.")
    add_number(doc, "En una terminal, abre la carpeta backend y ejecuta npm run dev o el script de inicio definido en package.json.")
    add_number(doc, "En otra terminal, abre la carpeta frontend y ejecuta npm start.")
    add_number(doc, "Abre http://localhost:3000 y realiza Ctrl + F5 si ves el nombre o el estilo antiguo.")
    add_heading(doc, "7.2 Comprobaciones rápidas", level=2)
    add_bullet(doc, "Visita http://localhost:5000/api/health. Debe aparecer status: ok y database: connected.")
    add_bullet(doc, "Visita http://localhost:5000/api/cliente?limit=100. Debes ver un arreglo de clientes en JSON.")
    add_bullet(doc, "En el frontend abre Clientes, pulsa Agregar, completa los campos obligatorios y Guarda. El registro debe aparecer en la tabla.")
    add_bullet(doc, "Edita un registro y confirma que cambia al guardar.")
    add_bullet(doc, "Pulsa Ocultar en un registro; debe desaparecer solo de esa tabla visual. Recarga la página para verificar el comportamiento guardado en Local Storage.")
    add_bullet(doc, "En el navegador presiona F12 > Aplicación > Local Storage > http://localhost:3000. Allí deben aparecer claves como venta-libros-cliente y venta-libros-historial después de usar un módulo.")

    add_heading(doc, "8. Errores comunes y solución")
    add_table(doc, ["Mensaje", "Causa", "Qué hacer"], [
        ("ECONNREFUSED", "MySQL no está iniciado o el puerto/host no coincide.", "Inicia MySQL en XAMPP y confirma la configuración del archivo backend/.env."),
        ("ER_NO_REFERENCED_ROW_2", "Se intentó guardar un ID relacionado que no existe.", "Selecciona una opción válida en las listas del formulario y guarda de nuevo."),
        ("No se puede eliminar porque el registro tiene relaciones activas", "El registro es usado por otra tabla.", "Ocúltalo si solo no quieres verlo; para borrarlo en MySQL elimina primero las relaciones dependientes."),
        ("No aparece lo guardado", "Se está viendo otra URL, datos sin recargar o MySQL detenido.", "Verifica /api/health, usa el frontend en puerto 3000 y actualiza con Ctrl + F5."),
        ("Aplicación de DevTools vacía", "No se ha generado aún almacenamiento local o no se seleccionó la clave.", "Abre un módulo, guarda/edita/oculta y selecciona una clave bajo Local Storage."),
    ], [1.75, 2.35, 2.4])

    add_heading(doc, "9. Guía de estudio para el compañero")
    add_number(doc, "Primero revisa la base de datos: identifica claves primarias, claves foráneas y por qué una venta necesita cliente y empleado.")
    add_number(doc, "Luego revisa el backend: index.js registra las rutas; cada ruta llama el controlador CRUD; db.js abre la conexión MySQL.")
    add_number(doc, "Después revisa el frontend: App.js define las rutas y tarjetas; modulo.js contiene el listado, formulario, carga de relaciones y Local Storage.")
    add_number(doc, "Prueba la API con el navegador antes de probar formularios. Así separas un problema de servidor de un problema visual.")
    add_number(doc, "Practica el ciclo completo: crear categoría, editorial, autor, libro, cliente, venta, detalle de venta y pago. Respeta ese orden porque hay relaciones entre tablas.")

    add_heading(doc, "10. Límites actuales y recomendaciones")
    add_bullet(doc, "La autenticación, permisos por rol y auditoría avanzada no forman parte de esta implementación; serían el siguiente paso para un entorno real.")
    add_bullet(doc, "El historial y las instantáneas de Local Storage no reemplazan a MySQL; se usan como apoyo visual y para inspección en el navegador.")
    add_bullet(doc, "Para borrar permanentemente, usa una operación DELETE consciente de las relaciones. Para ocultar solo de la página, usa Ocultar.")
    add_bullet(doc, "Antes de entregar el proyecto, conserva una copia del archivo SQL y del archivo .env (sin publicar contraseñas reales).")

    add_callout(doc, "Conclusión.", "La aplicación quedó conectada a la base libreria, con módulos CRUD, validación de relaciones, mensajes de error comprensibles y una interfaz React que permite estudiar y probar el flujo completo de ventas de libros.")
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
